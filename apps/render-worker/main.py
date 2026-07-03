from __future__ import annotations

import io
import os
from typing import Any

import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

app = FastAPI(title='hermes-render-worker', version='0.1.0')

# 中文图表标签：优先 Noto CJK（Docker 已装），本地 macOS 可回退 PingFang / DejaVu
plt.rcParams['font.sans-serif'] = [
    'Noto Sans CJK SC',
    'WenQuanYi Micro Hei',
    'PingFang SC',
    'Heiti SC',
    'Arial Unicode MS',
    'DejaVu Sans',
]
plt.rcParams['axes.unicode_minus'] = False


class RenderWordRequest(BaseModel):
    spec: dict[str, Any]
    echartsOptions: list[dict[str, Any]] | None = None


class RenderChartRequest(BaseModel):
    chartType: str
    chartConfig: dict[str, Any]
    rows: list[dict[str, Any]]


def _to_float(value: Any) -> float:
    try:
        if value is None or value == '':
            return 0.0
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def render_chart_png(req: RenderChartRequest) -> bytes:
    rows = req.rows
    if not rows:
        raise ValueError('no rows for chart')

    cfg = req.chartConfig
    x_field = cfg.get('xField') or list(rows[0].keys())[0]
    y_field = cfg.get('yField') or list(rows[0].keys())[-1]
    xs = [str(row.get(x_field, '')) for row in rows]
    ys = [_to_float(row.get(y_field)) for row in rows]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    title = cfg.get('title') or '报表图表'
    ax.set_title(str(title))
    chart_type = (req.chartType or 'line').lower()

    if chart_type == 'bar':
        ax.bar(xs, ys, color='#f97316')
    elif chart_type == 'pie':
        ax.pie(ys, labels=xs, autopct='%1.1f%%')
    elif chart_type == 'table':
        raise ValueError('table chart has no png')
    else:
        ax.plot(xs, ys, marker='o', color='#f97316', linewidth=2)

    ax.set_xlabel(str(x_field))
    ax.set_ylabel(str(y_field))
    ax.tick_params(axis='x', rotation=30)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    return buf.getvalue()


def _add_data_table(doc: Document, rows: list[dict[str, Any]], max_rows: int = 50) -> None:
    if not rows:
        doc.add_paragraph('（无数据行）')
        return

    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = str(header)

    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = '' if row.get(header) is None else str(row.get(header))

    if len(rows) > max_rows:
        doc.add_paragraph(f'（仅展示前 {max_rows} 行，共 {len(rows)} 行）')


@app.get('/health')
def health() -> dict[str, str]:
    return {'status': 'ok'}


@app.post('/render/chart')
def render_chart(req: RenderChartRequest) -> Response:
    try:
        png = render_chart_png(req)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return Response(content=png, media_type='image/png')


@app.post('/render/word')
def render_word(req: RenderWordRequest) -> Response:
    spec = req.spec
    narrative = spec.get('narrative') or {}
    data = spec.get('data') or {}
    rows: list[dict[str, Any]] = data.get('rows') or []
    charts: list[dict[str, Any]] = spec.get('charts') or []

    doc = Document()
    title = str(spec.get('title') or '数据报表')
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    user_query = spec.get('userQuery')
    if user_query:
        p = doc.add_paragraph()
        p.add_run('用户问题：').bold = True
        p.add_run(str(user_query))

    summary = str(narrative.get('summary') or '')
    if summary:
        doc.add_heading('分析摘要', level=1)
        doc.add_paragraph(summary)

    insights = narrative.get('insights') or []
    if insights:
        doc.add_heading('关键洞察', level=2)
        for insight in insights:
            doc.add_paragraph(str(insight), style='List Bullet')

    for section in narrative.get('sections') or []:
        doc.add_heading(str(section.get('title') or '章节'), level=2)
        body = section.get('body')
        if body:
            doc.add_paragraph(str(body))

    caveats = narrative.get('caveats') or []
    if caveats:
        doc.add_heading('说明与限制', level=2)
        for item in caveats:
            doc.add_paragraph(str(item), style='List Bullet')

    chart_added = False
    for chart in charts:
        chart_type = str(chart.get('chartType') or 'line').lower()
        if chart_type == 'table' or not rows:
            continue
        chart_cfg = chart.get('chartConfig') or {}
        try:
            png = render_chart_png(
                RenderChartRequest(chartType=chart_type, chartConfig=chart_cfg, rows=rows),
            )
        except ValueError:
            continue

        chart_title = chart_cfg.get('title') or f'{chart_type} 图'
        doc.add_heading(str(chart_title), level=2)
        stream = io.BytesIO(png)
        doc.add_picture(stream, width=Inches(6))
        chart_added = True

    if rows:
        doc.add_heading('数据明细', level=2)
        meta = doc.add_paragraph(
            f"共 {data.get('rowCount', len(rows))} 行"
            + ('（数据已截断）' if data.get('truncated') else ''),
        )
        meta.runs[0].font.size = Pt(10)
        _add_data_table(doc, rows)

    created_at = spec.get('createdAt')
    if created_at:
        footer = doc.add_paragraph(f'生成时间：{created_at}')
        footer.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()
    if len(content) < 100 or content[:2] != b'PK':
        raise HTTPException(status_code=500, detail='generated docx is invalid')

    return Response(
        content=content,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename="report.docx"'},
    )


if __name__ == '__main__':
    import uvicorn

    port = int(os.environ.get('PORT', '4060'))
    uvicorn.run(app, host='0.0.0.0', port=port)
